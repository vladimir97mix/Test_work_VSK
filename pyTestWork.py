import os
import time
import shutil
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from pyOpenRPA.Robot import UIDesktop

gChromeExeFullPath = r'Resources\GoogleChromePortable\App\Chrome-bin\chrome.exe'
gExtensionFullPathList = []
gWebDriverFullPath = r'Resources\SeleniumWebDrivers\Chrome\chromedriver_win32 v84.0.4147.30\chromedriver.exe'


def WebDriverInit(inWebDriverFullPath, inChromeExeFullPath, inExtensionFullPathList):
    # Определение полного пути к брузеру и драйверу
    lWebDriverChromeOptionsInstance = webdriver.ChromeOptions()
    lWebDriverChromeOptionsInstance.binary_location = inChromeExeFullPath
    # Добавление расширений
    for lExtensionItemFullPath in inExtensionFullPathList:
        lWebDriverChromeOptionsInstance.add_extension(lExtensionItemFullPath)
    # Запустить экземпляр Chrome
    lWebDriverInstance = None
    if inWebDriverFullPath:
        # Запустить с указанным путем веб-драйвера
        lWebDriverInstance = webdriver.Chrome(executable_path=inWebDriverFullPath,
                                              options=lWebDriverChromeOptionsInstance)
    else:
        lWebDriverInstance = webdriver.Chrome(options=lWebDriverChromeOptionsInstance)
    # Вернуть результат
    return lWebDriverInstance


def FindElemets(webDriver):
    webDriver.get('https://yandex.ru/')  # Переходим на нужную страницу
    wait = WebDriverWait(webDriver, 10)  # Таймаут ожидания
    # Дожидаемся до возможности кликнуть на строку поиска
    element = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, "input.input__control.input__input")))
    element.send_keys("Купить страховку")  # Отправляем запрос в поисковую строку
    element.send_keys(Keys.RETURN)  # Нажимаем на ENTER

    elementMain = webDriver.find_elements(By.CLASS_NAME, 'serp-item')  # Основные эелементы поиска (li)
    elementLink = webDriver.find_elements(By.CLASS_NAME, 'OrganicTitle-Link')  # Линки для перехода по ссылкам

    titleTextList = []  # Список заголовков
    extendetTextLis = []  # Список подзаголовков

    for elem in elementMain:
        # Ищем в элементах заголовки и вносим текст в список
        if elem.find_elements(By.CLASS_NAME, 'OrganicTitle-LinkText'):
            elem_text = elem.find_elements(By.CLASS_NAME, 'OrganicTitle-LinkText')
            titleTextList.append(elem_text[0].text)
        # Ищем в элементах позаголовки с классом Organic-Text и вносим текст в список
        if elem.find_elements(By.CLASS_NAME, 'Organic-Text'):
            elem_text = elem.find_elements(By.CLASS_NAME, 'Organic-Text')
            extendetTextLis.append(elem_text[0].text)
        # Ищем в элементах позаголовки с классом text-container и вносим текст в список
        if elem.find_elements(By.CLASS_NAME, 'text-container.typo.typo_text_m.typo_line_m.organic__text'):
            elem_text = elem.find_elements(By.CLASS_NAME, 'text-container.typo.typo_text_m.typo_line_m.organic__text')
            extendetTextLis.append(elem_text[0].text)

    # Создаем папку для скринов
    if not os.path.exists('screens'):
        os.makedirs('screens')
    picNum = 1
    # Открытие ссылок результатов
    for elem in elementLink:
        elem.click()
        handles = webDriver.window_handles  # Список открытых вкладок в браузере

        for tab in handles[1:]:
            webDriver.switch_to_window(tab)  # Переключаемся по каждой вкладке, кроме первой - yandex.ru
            webDriver.get_screenshot_as_file('screens/scr' + str(picNum) + '.png')  # Делаем скрин в папку 'screens'
            picNum += 1
            webDriver.close()
            webDriver.switch_to_window(handles[0])

    # webDriver.quit()  # Закрытие браузера

    picNum = 1  # Нумерация заголовка
    strNum = 1  # Нумерация заголовка

    wpSel = [
        {"title": "Документ - WordPad", "class_name": "WordPadClass", "backend": "win32"}]  # Сформировали UIO селектор
    lExistBool = UIDesktop.UIOSelector_Exist_Bool(inUIOSelector=wpSel)  # Проверить наличие окна по UIO селектору
    if not lExistBool:  # Проверить наличие окна wordpad
        os.system("write")  # Открыть wordpad
        time.sleep(2)
    else:  # Проверить, что окно wordpad не свернуто
        UIOWordPad = UIDesktop.UIOSelector_Get_UIO(inSpecificationList=wpSel)  # Получить UIO экземпляр
        if UIOWordPad.is_minimized():  # Проверить, что wordpad находится в свернутом виде
            UIOWordPad.restore()  # Восстановить окно wordpad из свернутого вида
        # Проверить наличие UI элемента по UIO селектору
    lWP_IsExistBool = UIDesktop.UIOSelector_Exist_Bool(inUIOSelector=[{"title": "Документ - WordPad",
                                                                       "class_name": "WordPadClass",
                                                                       "backend": "win32"},
                                                                      {"title": "Окно Rich Text", "rich_text": "",
                                                                       "ctrl_index": 4}])
    if lWP_IsExistBool:
        uioRichText = UIDesktop.UIOSelector_Get_UIO(inSpecificationList=[{"title": "Документ - WordPad",
                                                                  "class_name": "WordPadClass",
                                                                  "backend": "win32"},
                                                                 {"title": "Окно Rich Text", "rich_text": "",
                                                                  "ctrl_index": 4}])
    uioPasteImg = UIDesktop.UIOSelector_Get_UIO([{"class_name": "WordPadClass", "backend": "uia"},
                                               {"title": "Изображение", "depth_start": 11, "depth_end": 11}])

    for txtTitle, txtExtTitle in zip( titleTextList, extendetTextLis):
        uioRichText.type_keys('^b')
        uioRichText.type_keys(str(strNum) + '{SPACE}' + txtTitle, with_spaces=True)
        uioRichText.type_keys('{ENTER}')
        uioRichText.type_keys('^b')
        uioRichText.type_keys(txtExtTitle, with_spaces=True)
        uioRichText.type_keys('{ENTER}')
        uioPasteImg.click_input()
        if UIDesktop.UIOSelector_Exist_Bool([{"class_name": "WordPadClass", "backend": "uia"},
                                          {"title": "Имя файла:", "depth_start": 3, "depth_end": 3}]):
            uioFilename = UIDesktop.UIOSelector_Get_UIO([{"class_name": "WordPadClass", "backend": "uia"},
                                                         {"title": "Имя файла:", "depth_start": 3, "depth_end": 3}])
        uioFilename.type_keys(os.path.abspath('screens/scr' + str(picNum) + '.png'))
        uioFilename.type_keys('{ENTER}')
        uioRichText.type_keys('{ENTER}')
        strNum += 1
        picNum += 1

    uioRichText.type_keys('^s')
    UIDesktop.UIOSelector_Exist_Bool([{"class_name": "WordPadClass", "backend": "uia"},
                                      {"title": "Имя файла:", "depth_start": 4, "depth_end": 4}])
    uioFilename = UIDesktop.UIOSelector_Get_UIO([{"class_name": "WordPadClass", "backend": "uia"},
                                                 {"title": "Имя файла:", "depth_start": 4, "depth_end": 4}])
    uioFilename.type_keys(os.path.abspath('Отчет по ключевой фразе_{}.rtf'.format(int(time.time()))), with_spaces=True)
    uioFilename.type_keys('{ENTER}')

    webDriver.quit()
    ####################################################################################################################
    # doc = Document()
    # for txtTitle, txtExt in zip (titleTextList, extendetTextLis):  # Циклл по спискам заголовков
    #     p = doc.add_paragraph()  # Создаем новый абзац
    #     run = p.add_run(str(strNum) + ' ' + txtTitle)  # Добавляем текст
    #     run.bold = True  # Делаем текст жирным
    #     doc.add_paragraph(txtExt)  # Вставляем подзаголовок в новый абзац
    #     doc.add_picture('screens/scr' + str(picNum) + '.png', width=Inches(4.0)) # Вставляем скриншот, шириной 4 дюйма
    #     picNum -= 1  # Декримент скринов
    #     strNum += 1  # Инкримент номера заголовка
    #
    # doc.save('Отчет по ключевой фразе_{}.docx'.format(int(time.time())))  # Сохранение файла
    ####################################################################################################################

# Инициализация веб драйвера
inWebDriver = WebDriverInit(gWebDriverFullPath, gChromeExeFullPath, gExtensionFullPathList)

# Запуск задачи для веб драйвера
FindElemets(inWebDriver)

shutil.rmtree('screens/', ignore_errors=True)  # Удаление папки скринов за ненадобностью
